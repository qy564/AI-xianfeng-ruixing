import re
import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

input_dir = r"C:\Users\华硕\Documents\Codex\2026-07-18\al-agent-ai-agent-ai-agent\outputs\加分材料"
output_path = os.path.join(input_dir, "merged_doc.docx")

file_order = [
    "README.md",
    "01_参考资料清单.md",
    "02_数据分析样本.md",
    "03_用户旅程对比图.md",
    "04_竞品深度分析矩阵.md",
    "05_技术架构详图.md",
    "06_研究笔记.md",
    "07_Agent对话实录模拟.md",
]

doc = Document()

for section in doc.sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)

style = doc.styles["Normal"]
font = style.font
font.name = "Microsoft YaHei"
font.size = Pt(11)
style.element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")

def set_run_font(run, name="Microsoft YaHei", size=None, bold=False, color=None, italic=False):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    if size:
        run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)
    run.font.italic = italic

def add_bullet(text, level=0, bold_prefix=""):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Cm(1.27 * (level + 1))
    p.paragraph_format.space_after = Pt(2)
    if bold_prefix:
        r = p.add_run(bold_prefix)
        set_run_font(r, bold=True)
        r2 = p.add_run(text)
        set_run_font(r2)
    else:
        r = p.add_run(text)
        set_run_font(r)

def add_code_block(code_text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.space_before = Pt(4)
    run = p.add_run(code_text)
    run.font.name = "Consolas"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

def parse_inline_md(paragraph, text):
    parts = re.split(r"(\*\*.*?\*\*|\*.*?\*|`.*?`)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            r = paragraph.add_run(part[2:-2])
            set_run_font(r, bold=True)
        elif part.startswith("*") and part.endswith("*") and not part.startswith("**"):
            r = paragraph.add_run(part[1:-1])
            set_run_font(r, italic=True)
        elif part.startswith("`") and part.endswith("`"):
            r = paragraph.add_run(part[1:-1])
            r.font.name = "Consolas"
            r.font.size = Pt(9)
        else:
            r = paragraph.add_run(part)
            set_run_font(r)

def process_markdown_table(doc, lines):
    rows = []
    for line in lines:
        stripped = line.strip()
        if not stripped or stripped.startswith("|---"):
            continue
        cells = [c.strip() for c in stripped.split("|")[1:-1]]
        rows.append(cells)
    if len(rows) < 2:
        return
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, row_data in enumerate(rows):
        for j, cell_text in enumerate(row_data):
            cell = table.cell(i, j)
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(cell_text)
            if i == 0:
                set_run_font(run, bold=True, size=9)
            else:
                set_run_font(run, size=9)
    doc.add_paragraph()

for filename in file_order:
    filepath = os.path.join(input_dir, filename)
    if not os.path.exists(filepath):
        continue
    with open(filepath, "r", encoding="utf-8") as f:
        content = f.read()
    lines = content.split("\n")
    i = 0
    in_code_block = False
    code_buffer = []
    in_table = False
    table_buffer = []

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if stripped.startswith("```"):
            if in_code_block:
                add_code_block("\n".join(code_buffer))
                code_buffer = []
                in_code_block = False
                i += 1
                continue
            else:
                in_code_block = True
                i += 1
                continue
        if in_code_block:
            code_buffer.append(line)
            i += 1
            continue

        if stripped.startswith("|") and stripped.endswith("|"):
            table_buffer.append(line)
            in_table = True
            i += 1
            continue
        else:
            if in_table:
                process_markdown_table(doc, table_buffer)
                table_buffer = []
                in_table = False

        if not stripped:
            i += 1
            continue

        if stripped.startswith("# "):
            doc.add_paragraph()
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(6)
            run = p.add_run(stripped[2:])
            set_run_font(run, size=18, bold=True)
            i += 1
            continue
        elif stripped.startswith("## "):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(10)
            p.paragraph_format.space_after = Pt(4)
            run = p.add_run(stripped[3:])
            set_run_font(run, size=15, bold=True)
            i += 1
            continue
        elif stripped.startswith("### "):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after = Pt(4)
            run = p.add_run(stripped[4:])
            set_run_font(run, size=13, bold=True)
            i += 1
            continue
        elif stripped.startswith("#### "):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after = Pt(4)
            run = p.add_run(stripped[5:])
            set_run_font(run, size=12, bold=True)
            i += 1
            continue

        if stripped in ("---", "***", "___"):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
            run = p.add_run("─" * 50)
            set_run_font(run, size=8, color=(180, 180, 180))
            i += 1
            continue

        if stripped.startswith("- ") or stripped.startswith("* ") or stripped.startswith("+ "):
            content_text = stripped[2:]
            m = re.match(r"\*\*(.*?)\*\*\s*(.*)", content_text)
            if m:
                bold_part = m.group(1) + " "
                rest = m.group(2)
                add_bullet(rest, bold_prefix=bold_part)
            else:
                add_bullet(content_text)
            i += 1
            continue

        if re.match(r"^\d+[\.\、]", stripped):
            content_text = re.sub(r"^\d+[\.\、]\s*", "", stripped)
            p = doc.add_paragraph(style="List Number")
            p.paragraph_format.space_after = Pt(2)
            parse_inline_md(p, content_text)
            i += 1
            continue

        if stripped:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(4)
            parse_inline_md(p, stripped)

        i += 1

    if in_table:
        process_markdown_table(doc, table_buffer)
        table_buffer = []
        in_table = False

    if filename != file_order[-1]:
        doc.add_page_break()

doc.save(output_path)
print("Done: " + output_path)
